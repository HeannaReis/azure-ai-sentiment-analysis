# services/document_service.py
from docx import Document
from docx.shared import Inches
from pathlib import Path
from core.logger_config import logger
from config import Config

class DocumentService:
    """Serviço para geração de documentos Word."""
    
    def __init__(self):
        """Inicializa o serviço de documento."""
        self.document = Document()
        self.document.add_heading('Resumo de Análises de Imagens', 0)
        
    def add_image_summary(self, image_name: str, summary: str):
        """
        Adiciona um resumo de imagem ao documento.
        
        Args:
            image_name: Nome do arquivo de imagem
            summary: Texto do resumo
        """
        try:
            self.document.add_heading(f'Imagem: {image_name}', level=1)
            
            # Adiciona a imagem ao documento
            image_path = Config.PROCESSED_DIR / image_name
            if image_path.exists():
                self.document.add_picture(str(image_path), width=Inches(4))
            
            # Adiciona o resumo
            self.document.add_heading('Análise:', level=2)
            self.document.add_paragraph(summary)
            self.document.add_paragraph('')  # Espaçamento
            
            logger.info(f"Resumo da imagem '{image_name}' adicionado ao documento.")
        except Exception as e:
            logger.error(f"Erro ao adicionar resumo ao documento: {e}")
            
    def save_document(self):
        """Salva o documento no caminho configurado."""
        try:
            self.document.save(Config.OUTPUT_DOCX)
            logger.info(f"Documento salvo em {Config.OUTPUT_DOCX}")
        except Exception as e:
            logger.error(f"Erro ao salvar documento: {e}")